"""Markdown → DOCX renderer using python-docx.

Produces an editable Word document from the same parsed blocks that the PDF
renderer consumes. Brand colors are applied to headings and table headers;
full visual fidelity to the PDF is not a goal — editability is.
"""

from __future__ import annotations

import io
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .config import BrandConfig, LayoutOptions
from .parser import parse_markdown


_INLINE_RE = re.compile(
    r"(\[[^\]]+\]\([^)\s]+\)|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|⟦F\d+⟧)"
)

# Captures the label and target of a `[text](url)` inline link.
_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)\s]+)\)")

# Word's conventional hyperlink colour (blue), applied so links read as links.
_LINK_COLOR = "0563C1"

_HYPERLINK_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)


def _rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _add_hyperlink(para, url: str, text: str, color: str = _LINK_COLOR) -> None:
    """Append a clickable, coloured+underlined hyperlink run to `para`."""
    r_id = para.part.relate_to(url, _HYPERLINK_REL, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    rpr.append(color_el)
    u_el = OxmlElement("w:u")
    u_el.set(qn("w:val"), "single")
    rpr.append(u_el)
    run.append(rpr)

    t_el = OxmlElement("w:t")
    t_el.set(qn("xml:space"), "preserve")
    t_el.text = text
    run.append(t_el)

    hyperlink.append(run)
    para._p.append(hyperlink)


def _apply_inline(para, text: str) -> None:
    """Add runs to `para` with link/bold/italic/code formatting from inline markdown."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        start, end = m.span()
        if start > pos:
            para.add_run(text[pos:start])
        chunk = m.group(0)
        if chunk.startswith("["):
            link = _LINK_RE.match(chunk)
            _add_hyperlink(para, link.group(2), link.group(1))
        elif chunk.startswith("**"):
            para.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith("*"):
            para.add_run(chunk[1:-1]).italic = True
        elif chunk.startswith("`"):
            run = para.add_run(chunk[1:-1])
            run.font.name = "Courier New"
        elif chunk.startswith("⟦"):
            para.add_run(f"[{chunk[1:-1]}]")
        pos = end
    if pos < len(text):
        para.add_run(text[pos:])


def _configure_styles(doc, brand: BrandConfig, layout: LayoutOptions) -> None:
    """Apply the brand font family, layout sizes and colors to the base
    document styles so every block (paragraphs, lists, tables, quotes) inherits
    the brand look instead of Word's generic Calibri defaults.

    Note: the .docx references the font by name; it is not embedded. It renders
    with the brand font on machines that have it installed, otherwise Word
    substitutes a fallback.
    """
    family = brand.fonts.family

    normal = doc.styles["Normal"]
    normal.font.name = family
    normal.font.size = Pt(layout.body_size)
    normal.font.color.rgb = _rgb(brand.colors.text)

    _heading = brand.colors.heading
    headings = [
        ("Heading 1", layout.h1_size, _heading or brand.colors.primary_dark),
        ("Heading 2", layout.h2_size, _heading or brand.colors.primary_dark),
        ("Heading 3", layout.h3_size, _heading or brand.colors.primary_mid),
    ]
    for name, size, color in headings:
        style = doc.styles[name]
        style.font.name = family
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = _rgb(color)


def _set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#").upper())
    tcPr.append(shd)


def _set_para_shading(para, hex_color: str) -> None:
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#").upper())
    pPr.append(shd)


def _repeat_table_header(row) -> None:
    """Mark a row as a table header so Word repeats it at the top of every page
    the table spans (the DOCX equivalent of the PDF's header-repeat-on-break)."""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def _row_cant_split(row) -> None:
    """Prevent a single row's content from splitting across a page break."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def _add_hr(doc: Document) -> None:
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_footer_logo(doc: Document, brand: BrandConfig) -> None:
    """Place the brand logo in the section footer (bottom-left) — or the logo
    fallback text when the image is missing/unreadable — and the client
    co-branding logo bottom-right, mirroring the PDF."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    if brand.logo_path and Path(brand.logo_path).exists():
        try:
            para.add_run().add_picture(str(brand.logo_path), height=Pt(18))
        except Exception:
            _add_footer_fallback_text(para, brand)
    else:
        _add_footer_fallback_text(para, brand)

    if brand.client_logo_path and Path(brand.client_logo_path).exists():
        try:
            content_width = section.page_width - section.left_margin - section.right_margin
            para.paragraph_format.tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)
            para.add_run("\t")
            para.add_run().add_picture(str(brand.client_logo_path), height=Pt(18))
        except Exception:
            pass


def _add_footer_fallback_text(para, brand: BrandConfig) -> None:
    if brand.logo_fallback_text:
        run = para.add_run(brand.logo_fallback_text)
        run.font.size = Pt(8)
        run.font.color.rgb = _rgb(brand.colors.primary_dark)


def render_markdown_to_docx(
    md_text: str,
    brand: BrandConfig,
    layout: LayoutOptions | None = None,
    base_path: str | Path | None = None,
) -> bytes:
    """Render Markdown (with optional YAML frontmatter) to DOCX bytes."""
    layout = layout or LayoutOptions()
    meta, blocks = parse_markdown(md_text)

    doc = Document()
    _configure_styles(doc, brand, layout)

    # Page margins
    for section in doc.sections:
        section.left_margin = Pt(layout.margin_left)
        section.right_margin = Pt(layout.margin_right)
        section.top_margin = Pt(layout.margin_top)
        section.bottom_margin = Pt(layout.margin_bottom)

    content_width_pt = layout.page_size[0] - layout.margin_left - layout.margin_right
    content_width_in = content_width_pt / 72.0

    # Brand logo in the page footer (bottom-left), mirroring the PDF. Falls back
    # to the logo text when no image is available.
    _add_footer_logo(doc, brand)

    base = Path(base_path) if base_path else None

    for block in blocks:
        btype = block["type"]

        # Heading colors come from the configured styles (see _configure_styles),
        # so the runs only need their text.
        if btype in ("h1_doc", "h1"):
            _apply_inline(doc.add_heading("", level=1), block["text"])

        elif btype == "h2":
            _apply_inline(doc.add_heading("", level=2), block["text"])

        elif btype == "h3":
            _apply_inline(doc.add_heading("", level=3), block["text"])

        elif btype == "p":
            para = doc.add_paragraph()
            if layout.justify:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _apply_inline(para, block["text"])

        elif btype == "blockquote":
            for line in block["lines"]:
                para = doc.add_paragraph(style="Quote")
                _apply_inline(para, line)

        elif btype == "callout":
            label_para = doc.add_paragraph()
            label_run = label_para.add_run(block["label"])
            label_run.bold = True
            label_para.paragraph_format.left_indent = Pt(20)
            for line in block["lines"]:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Pt(20)
                _apply_inline(para, line)

        elif btype == "checklist":
            for item in block["items"]:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Pt(18)
                para.paragraph_format.first_line_indent = Pt(-18)
                ok = item.get("ok", True)
                marker = para.add_run(("✓ " if ok else "✗ "))
                marker.bold = True
                marker.font.color.rgb = (
                    _rgb(brand.colors.primary_light) if ok else RGBColor(0xD6, 0x45, 0x41)
                )
                _apply_inline(para, item["text"])

        elif btype == "ul":
            for item in block["items"]:
                para = doc.add_paragraph(style="List Bullet")
                _apply_inline(para, item["text"])
                for child in item.get("children", []):
                    cpara = doc.add_paragraph(style="List Bullet 2")
                    _apply_inline(cpara, child)

        elif btype == "ol":
            for item in block["items"]:
                para = doc.add_paragraph(style="List Number")
                _apply_inline(para, item)

        elif btype == "table":
            header = block["header"]
            rows = block["rows"]
            n_cols = len(header)
            tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
            tbl.style = "Table Grid"

            # Header row: repeat at the top of every page the table spans, and
            # keep it glued to the first data row so it never sits orphaned at
            # the bottom of a page.
            hdr_row = tbl.rows[0]
            _repeat_table_header(hdr_row)
            _row_cant_split(hdr_row)
            for j, cell_text in enumerate(header):
                cell = hdr_row.cells[j]
                _set_cell_shading(cell, brand.colors.table_header_bg or brand.colors.primary_dark)
                para = cell.paragraphs[0]
                para.paragraph_format.keep_with_next = True
                run = para.add_run(cell_text)
                run.bold = True
                run.font.color.rgb = _rgb(brand.colors.white)

            # Data rows
            for ridx, row in enumerate(rows):
                tbl_row = tbl.rows[ridx + 1]
                _row_cant_split(tbl_row)
                # Shade the whole row (all grid cells) so ragged rows with fewer
                # cells than the header don't show half-striped backgrounds.
                if ridx % 2 == 1:
                    for cell in tbl_row.cells:
                        _set_cell_shading(cell, brand.colors.bg_soft)
                for j, cell_text in enumerate(row[:n_cols]):
                    _apply_inline(tbl_row.cells[j].paragraphs[0], cell_text)

        elif btype == "code":
            if not block["text"].strip():
                continue
            lines = block["text"].splitlines()
            for i, line in enumerate(lines):
                para = doc.add_paragraph(style="No Spacing")
                _set_para_shading(para, brand.colors.bg_soft)
                run = para.add_run(line if line else " ")
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                if i == 0 and block.get("lang"):
                    para.paragraph_format.space_before = Pt(6)
                if i == len(lines) - 1:
                    para.paragraph_format.space_after = Pt(6)

        elif btype == "image":
            # Absolute paths resolve on their own (e.g. mermaid/formula PNGs the
            # web layer stages into the FS); relative ones need `base`.
            src = block["src"]
            if Path(src).is_absolute():
                img_path = Path(src)
            elif base:
                img_path = base / src
            else:
                img_path = None
            if img_path is not None and img_path.exists():
                # python-docx raises on unrecognized/corrupt images (e.g. an
                # SVG or a truncated PNG). Skip rather than abort the render.
                try:
                    doc.add_picture(str(img_path), width=Inches(content_width_in))
                except Exception:
                    pass

        elif btype == "formula":
            para = doc.add_paragraph()
            run = para.add_run(f"[Formula: {block['id']}]")
            run.italic = True

        elif btype == "quiz":
            q_para = doc.add_paragraph()
            _apply_inline(q_para, block["question"])
            for opt in block["options"]:
                o_para = doc.add_paragraph(style="List Bullet")
                _apply_inline(o_para, opt)

        elif btype == "hr":
            _add_hr(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
